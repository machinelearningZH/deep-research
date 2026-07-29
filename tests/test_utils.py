from __future__ import annotations

import io

import pytest
from docx import Document

from _core import utils


@pytest.mark.parametrize(
    ("development_enabled", "fast_mode", "expected_models", "expected_workflow"),
    [
        (
            True,
            False,
            {"query": "development", "report": "development"},
            {"max_queries": 2, "search_limit": 1, "auto_limit": 1},
        ),
        (
            False,
            True,
            {"query": "low", "report": "low"},
            {"max_queries": 5, "search_limit": 4, "auto_limit": 2},
        ),
        (
            False,
            False,
            {"query": "low", "report": "high"},
            {"max_queries": 20, "search_limit": 15, "auto_limit": 3},
        ),
    ],
    ids=["development-overrides-fast", "fast", "standard"],
)
def test_get_model_and_workflow_config_selects_complete_mode(
    monkeypatch: pytest.MonkeyPatch,
    development_enabled: bool,
    fast_mode: bool,
    expected_models: dict[str, str],
    expected_workflow: dict[str, int],
) -> None:
    config = {
        "development": {
            "enabled": development_enabled,
            "model_for_dev_testing": "development",
        },
        "models": {
            "performance_low": "low",
            "performance_medium": "medium",
            "performance_high": "high",
        },
        "app": {
            "max_queries_dev": 2,
            "search_limit_dev": 1,
            "search_auto_limit_dev": 1,
            "max_queries_fast": 5,
            "search_limit_fast": 4,
            "search_auto_limit_fast": 2,
            "max_queries": 20,
            "search_limit": 15,
            "search_auto_limit": 3,
        },
    }
    monkeypatch.setattr(utils, "config", config)

    models, workflow = utils.get_model_and_workflow_config(fast_mode=fast_mode)

    assert models["create_queries"] == expected_models["query"]
    assert models["final_report"] == expected_models["report"]
    assert set(models) == {
        "create_queries",
        "check_relevance",
        "analyze_documents",
        "reflect_task",
        "final_report",
    }
    assert workflow == expected_workflow


def test_parallel_calls_preserve_order_and_isolate_failures() -> None:
    def process(prompt: str, *, suffix: str) -> str:
        if prompt == "broken":
            raise ValueError("service unavailable")
        return f"{prompt}{suffix}"

    results = utils.call_function_in_parallel(
        ["first", "broken", "third"],
        process,
        max_workers=3,
        suffix="-done",
    )

    assert results == ["first-done", "Error: service unavailable", "third-done"]


def test_create_docx_from_markdown_preserves_document_semantics() -> None:
    markdown = """## Result
Text with **bold**, *italic*, and [source](https://example.com).
- First item
  - Nested item
"""

    document_bytes = utils.create_docx_from_markdown("Why?", markdown)
    document = Document(io.BytesIO(document_bytes))

    assert document.paragraphs[0].text == "Recherche-Bericht"
    assert "Recherchefrage: Why?" in [
        paragraph.text for paragraph in document.paragraphs
    ]
    assert any(
        run.text == "bold" and run.bold
        for paragraph in document.paragraphs
        for run in paragraph.runs
    )
    assert any(
        run.text == "italic" and run.italic
        for paragraph in document.paragraphs
        for run in paragraph.runs
    )
    assert [paragraph.style.name for paragraph in document.paragraphs][-2:] == [
        "List Bullet",
        "List Bullet 2",
    ]
    assert any(
        relationship.target_ref == "https://example.com"
        for relationship in document.part.rels.values()
    )
